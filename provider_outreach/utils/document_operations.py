#!/usr/bin/env python3
"""
Document Operations Module
Handles document generation and manipulation for provider outreach
"""

import logging
from pathlib import Path
from typing import Dict, List
from datetime import datetime
from docx import Document
import PyPDF2

logger = logging.getLogger(__name__)

# Constants for document processing
ACCEPTABLE_MODIFIERS = ["26", "TC", "59", "76", "77", "91", "99"]
ACCEPTABLE_POS = ["11", "22", "23", "24", "31", "32", "33", "34", "41", "42", "49", "50", "51", "52", "53", "54", "55", "56", "57", "58", "59", "60", "61", "62", "63", "64", "65", "71", "72", "81", "99"]

class DocumentOperations:
    def __init__(self, template_path: Path):
        """Initialize with path to Word template"""
        self.template_path = template_path
        
    def process_line_items(self, line_items: List[Dict]) -> Dict:
        """Process line items for document placeholders"""
        mapping = {}
        for i, line in enumerate(line_items[:6], start=1):
            # Get modifier, handling both string format and list format
            modifier_raw = line.get("modifiers", [])
            if isinstance(modifier_raw, list):
                modifier = ",".join([m for m in modifier_raw if m in ACCEPTABLE_MODIFIERS])
            elif isinstance(modifier_raw, str):
                modifier = modifier_raw if modifier_raw in ACCEPTABLE_MODIFIERS else ""
            else:
                modifier = ""
            
            pos = line.get("place_of_service", "11")  # Default POS
            units = line.get("units", 1)
            charge = float(line.get("charge_amount", 0))
            
            mapping.update({
                f"<dos{i}>": line.get("date_of_service", ""),
                f"<cpt{i}>": line.get("cpt_code", "N/A"),
                f"<charge{i}>": "${:,.2f}".format(charge),
                f"<units{i}>": units,
                f"<modifier{i}>": modifier,
                f"<pos{i}>": pos,
                f"<alwd{i}>": "${:,.2f}".format(charge),  # Using charge as allowed amount for now
                f"<paid{i}>": "${:,.2f}".format(charge),  # Using charge as paid amount for now
                f"<code{i}>": "CO-50"  # CO-50 denial code
            })
            
        # Fill in empty values for any remaining rows
        for i in range(len(line_items) + 1, 7):
            mapping.update({
                f"<dos{i}>": "", f"<cpt{i}>": "", f"<charge{i}>": "", f"<units{i}>": "",
                f"<modifier{i}>": "", f"<pos{i}>": "", f"<alwd{i}>": "", f"<paid{i}>": "", f"<code{i}>": ""
            })
            
        return mapping
        
    def populate_placeholders(self, doc: Document, mapping: Dict):
        """Replace placeholders in a Word document with values"""
        sanitized_mapping = {k: str(v) if v is not None else "" for k, v in mapping.items()}
        
        # Process paragraphs
        for paragraph in doc.paragraphs:
            for placeholder, value in sanitized_mapping.items():
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, value)
                    
        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for placeholder, value in sanitized_mapping.items():
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, value)
                            
    def generate_outreach_document(self, json_data: Dict, db_data: Dict, 
                                 output_path: Path) -> bool:
        """Generate outreach document from template"""
        try:
            # Load the template
            doc = Document(self.template_path)
            
            # Basic document info - primarily from DB, with patient account number from JSON
            mapping = {
                "<process_date>": datetime.now().strftime("%Y-%m-%d"),
                "<PatientName>": db_data.get("PatientName", "N/A"),
                "<dob>": db_data.get("Patient_DOB", ""),
                "<doi>": db_data.get("Patient_Injury_Date", ""),
                "<provider_ref>": json_data.get("billing_info", {}).get("patient_account_no", "N/A"),  # From JSON
                "<order_no>": db_data.get("FileMaker_Record_Number", "N/A"),
                "<billing_name>": db_data.get("Billing Name", "N/A"),
                "<billing_address1>": db_data.get("Billing Address 1", "N/A"),
                "<billing_address2>": db_data.get("Billing Address 2", ""),
                "<billing_city>": db_data.get("Billing Address City", "N/A"),
                "<billing_state>": db_data.get("Billing Address State", "N/A"),
                "<billing_zip>": db_data.get("Billing Address Postal Code", "N/A"),
                "<TIN>": db_data.get("TIN", "N/A"),
                "<NPI>": db_data.get("NPI", "N/A"),
                "<total_paid>": "${:,.2f}".format(float(json_data.get("billing_info", {}).get("total_charge", 0))),  # From JSON
            }
            
            # Add line item details from JSON service_lines
            mapping.update(self.process_line_items(json_data.get("service_lines", [])))
            
            # Populate placeholders
            self.populate_placeholders(doc, mapping)
            
            # Save the document
            doc.save(output_path)
            return True
            
        except Exception as e:
            logger.error(f"Error generating document: {str(e)}")
            return False
            
    def merge_pdfs(self, pdf1_path: Path, pdf2_path: Path, 
                  output_path: Path) -> bool:
        """Merge two PDF files"""
        try:
            merger = PyPDF2.PdfMerger()
            
            # Add the first PDF
            merger.append(str(pdf1_path))
            
            # Add the second PDF
            merger.append(str(pdf2_path))
            
            # Write the merged PDF
            merger.write(str(output_path))
            merger.close()
            
            return True
            
        except Exception as e:
            logger.error(f"Error merging PDFs: {str(e)}")
            return False 